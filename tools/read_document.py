"""read_document — Doc §4.2 agent tool.

Reads an uploaded file (PDF, Word, Excel) and returns extracted text + metadata.
"""

from __future__ import annotations

import json
from pathlib import Path


def read_document(file_path: str) -> str:
    """Extract text from a local PDF/Word/Excel file."""
    path = Path(file_path)
    if not path.exists():
        return json.dumps({"error": f"File not found: {file_path}"})

    suffix = path.suffix.lower()
    try:
        if suffix == ".pdf":
            from pypdf import PdfReader
            reader = PdfReader(path)
            text = "\n".join(p.extract_text() or "" for p in reader.pages)
            meta = {"pages": len(reader.pages)}
        elif suffix in (".docx", ".doc"):
            from docx import Document
            doc = Document(path)
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            meta = {"paragraphs": len(doc.paragraphs)}
        elif suffix in (".xlsx", ".xls"):
            from openpyxl import load_workbook
            wb = load_workbook(path, data_only=True)
            parts = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                parts.append(f"## Sheet: {sheet_name}")
                for row in ws.iter_rows(values_only=True):
                    parts.append(" | ".join(str(c) if c is not None else "" for c in row))
            text = "\n".join(parts)
            meta = {"sheets": wb.sheetnames}
        else:
            text = path.read_text(errors="replace")
            meta = {"encoding": "utf-8"}
    except Exception as e:
        return json.dumps({"error": str(e), "file_path": file_path})

    return json.dumps({
        "file_path": file_path,
        "metadata": meta,
        "content": text[:20_000],
        "truncated": len(text) > 20_000,
    })
